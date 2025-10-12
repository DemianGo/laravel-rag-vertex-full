#!/usr/bin/env python3
"""
DOCX Extractor Wrapper
Directly extracts DOCX files, works with temp files without extensions.
"""

import sys
import os
import shutil

def extract_docx(file_path: str) -> str:
    """Extract text from DOCX"""
    temp_link = None
    try:
        if not os.path.exists(file_path):
            return ''
        
        # python-docx may work better with .docx extension
        if not file_path.endswith('.docx') and not file_path.endswith('.doc'):
            temp_link = file_path + '.docx'
            shutil.copy2(file_path, temp_link)
            file_to_read = temp_link
        else:
            file_to_read = file_path
        
        from docx import Document
        
        doc = Document(file_to_read)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return '\n'.join(text_parts)
    except Exception as e:
        # Silent fail - return empty
        return ''
    finally:
        # Clean up temp file
        if temp_link and os.path.exists(temp_link):
            try:
                os.remove(temp_link)
            except:
                pass

if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit(1)
    
    file_path = sys.argv[1]
    text = extract_docx(file_path)
    print(text)

