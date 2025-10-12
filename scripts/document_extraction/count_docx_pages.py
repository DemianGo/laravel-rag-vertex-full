#!/usr/bin/env python3
import sys
import os
import tempfile
import shutil
try:
    from docx import Document
    input_file = sys.argv[1]
    
    # Handle files without extension
    if not input_file.endswith('.docx'):
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_file.close()
        shutil.copy2(input_file, temp_file.name)
        input_file = temp_file.name
        delete_temp = True
    else:
        delete_temp = False
    
    doc = Document(input_file)
    
    # Count page breaks + estimate from paragraphs
    page_breaks = sum(1 for p in doc.paragraphs if '\\f' in p.text or '\\x0c' in p.text)
    total_paragraphs = len(doc.paragraphs)
    
    # Estimate: 30 paragraphs per page (average)
    estimated_pages = max(page_breaks, total_paragraphs // 30)
    print(max(1, estimated_pages))
    
    if delete_temp and os.path.exists(input_file):
        os.unlink(input_file)
except Exception as e:
    sys.exit(1)